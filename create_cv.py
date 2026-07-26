"""Generate Fredrick Mahinay's master CV as a .docx file."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style defaults ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(4)
pf.line_spacing = 1.15

# ── Helper functions ──

def add_horizontal_line(doc):
    """Add a thin horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_heading_styled(doc, text, level=1):
    """Add a section heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        run.font.name = 'Calibri'
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    add_horizontal_line(doc)
    return h


def add_subheading(doc, text):
    """Add a job title / subsection heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_meta_line(doc, text):
    """Add company / date / location metadata line."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text):
    """Add a bullet-point paragraph."""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'
    return p


def add_body(doc, text):
    """Add a normal body paragraph."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_project_block(doc, title, status, stack, links_text, bullets):
    """Add a project entry."""
    add_subheading(doc, title)
    p = doc.add_paragraph()
    run = p.add_run(f"Status: {status}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(1)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Tech Stack: {stack}")
    run2.font.size = Pt(10)
    run2.font.name = 'Calibri'
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p2.paragraph_format.space_after = Pt(1)

    p3 = doc.add_paragraph()
    run3 = p3.add_run(links_text)
    run3.font.size = Pt(10)
    run3.font.name = 'Calibri'
    run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p3.paragraph_format.space_after = Pt(4)

    for b in bullets:
        add_bullet(doc, b)


# ══════════════════════════════════════════════════════════════
# HEADER — Name & contact
# ══════════════════════════════════════════════════════════════

name_para = doc.add_paragraph()
name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = name_para.add_run('FREDRICK MAHINAY')
name_run.bold = True
name_run.font.size = Pt(22)
name_run.font.name = 'Calibri'
name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
name_para.paragraph_format.space_after = Pt(2)

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag_run = tagline.add_run('Civil Engineer \u2192 Network Systems Specialist \u2192 AI Application Builder')
tag_run.font.size = Pt(11)
tag_run.font.name = 'Calibri'
tag_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
tagline.paragraph_format.space_after = Pt(6)

# Contact details — single centered line
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_run = contact.add_run(
    'Bi\u00f1an, Laguna, Philippines  |  '
    'fredrickmahinay18@gmail.com  |  '
    '0961 003 0671'
)
contact_run.font.size = Pt(9.5)
contact_run.font.name = 'Calibri'
contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
contact.paragraph_format.space_after = Pt(2)

links = doc.add_paragraph()
links.alignment = WD_ALIGN_PARAGRAPH.CENTER
links_run = links.add_run(
    'Portfolio: fredrickmahinay.github.io  |  '
    'GitHub: github.com/burny143  |  '
    'LinkedIn: linkedin.com/in/fredrick-mahinay-62a4a8358'
)
links_run.font.size = Pt(9.5)
links_run.font.name = 'Calibri'
links_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
links.paragraph_format.space_after = Pt(6)

add_horizontal_line(doc)

# ══════════════════════════════════════════════════════════════
# PROFESSIONAL SUMMARY
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, 'Professional Summary', level=2)

summary = (
    'Licensed Civil Engineer with six years of hands-on construction and structural QA/QC experience, '
    'who successfully transitioned into enterprise IT and is now building toward a full-time career in '
    'AI engineering. Since 2023, has applied the same inspect-before-you-trust discipline learned on '
    'live construction sites to enterprise network operations, supporting infrastructure at Accenture '
    'and managing Tier-2 network and security operations for the Asian Development Bank project at Indra. '
    'Currently self-directing a transition into applied AI development, treating AI as an active collaborator '
    'across architecture decisions, data pipelines, and interface design while independently designing, '
    'building, and shipping two full-stack applications end to end: a hero-matchup analytics engine for '
    'Mobile Legends: Bang Bang and a live cryptocurrency market dashboard with its own automated ETL pipeline. '
    'Combines a rare cross-disciplinary foundation, rigorous field-tested quality assurance habits, enterprise '
    'systems troubleshooting, and hands-on software delivery, with a fast-growing, self-taught skill set in '
    'Python, SQL, Supabase, and AI-assisted development workflows. Seeking to move into an AI engineering, '
    'AI-assisted application development, or systems integration role where this combination of discipline, '
    'adaptability, and builder\u2019s mindset can be applied at scale.'
)
add_body(doc, summary)

# Key highlights
highlights = [
    '6 years of civil engineering foundation across structural QA/QC, cost estimation, and project management',
    '2 AI-assisted applications independently designed, built, and shipped',
    '1,000+ end users supported across enterprise IT operations',
    'Licensed Civil Engineer (PRC License No. 0147997) and certified Dell Boomi Integration Developer',
]
p_hl = doc.add_paragraph()
p_hl.paragraph_format.space_before = Pt(4)
run_hl = p_hl.add_run('Key Highlights:')
run_hl.bold = True
run_hl.font.size = Pt(10.5)
run_hl.font.name = 'Calibri'
for h in highlights:
    add_bullet(doc, h)

# ══════════════════════════════════════════════════════════════
# TECHNICAL SKILLS
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, 'Technical Skills', level=2)

skill_categories = [
    ('Languages & Data', [
        'Python (ETL, data pipelines, scripting)',
        'SQL (Oracle SQL, Microsoft SQL Server)',
        'Excel / VBA (multi-sheet scoring models, dependent rule logic)',
        'JavaScript',
    ]),
    ('Frameworks, Platforms & Integration', [
        'Supabase (database, backend services)',
        'Dell Boomi (integration platform, Associate & Professional certified)',
        'TIBCO (enterprise integration troubleshooting)',
        'yfinance (financial market data ingestion)',
    ]),
    ('AI / ML & Prompting', [
        'Prompt engineering and iterative refinement',
        'AI-assisted software architecture and debugging',
        'Agentic workflows and automation',
    ]),
    ('Tools & Deployment', [
        'Git / GitHub & GitHub Pages',
        'GitHub Actions (CI/CD, automated data pipelines)',
        'ServiceNow (incident and change management)',
        'Active Directory',
        'Windows Server administration',
        'Zscaler (ZPA, Zscaler Private Access) \u2014 Tier 2 support',
        'MS Project (Gantt charting, timeline tracking)',
        '2D CAD drafting',
    ]),
    ('Network & Systems Support', [
        'L2 network escalation and root-cause incident resolution',
        'DNS troubleshooting and website whitelisting',
        'Wi-Fi and internet connectivity diagnostics',
        'Change Request (CHG) and Service Request (RITM) execution',
        'Incident Coordination (P1 to P3 severity levels)',
    ]),
    ('Civil & Technical Foundations', [
        'Structural QA/QC and site inspection',
        'Quantity take-off and cost estimation',
        'Contractor bidding, evaluation, and subcontractor management',
        'Site field surveys and 2D CAD shop drawings',
    ]),
]

for category, items in skill_categories:
    p_cat = doc.add_paragraph()
    run_cat = p_cat.add_run(category)
    run_cat.bold = True
    run_cat.font.size = Pt(10.5)
    run_cat.font.name = 'Calibri'
    run_cat.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p_cat.paragraph_format.space_before = Pt(6)
    p_cat.paragraph_format.space_after = Pt(2)
    for item in items:
        add_bullet(doc, item)

# ══════════════════════════════════════════════════════════════
# PROJECTS
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, 'Projects', level=2)

add_project_block(
    doc,
    title='MLBB Hero Counter-Pick Analyzer',
    status='In active development',
    stack='Excel / VBA, Python, Supabase, GitHub Pages',
    links_text='Demo: burny143.github.io/mlbb-drafter-v2.0/  |  Code: github.com/burny143/mlbb-drafter-v2.0',
    bullets=[
        'Designed and built a custom scoring engine that recommends the optimal counter-pick for any Mobile Legends: Bang Bang hero matchup, weighing seven distinct evaluation factors: role advantage, stat differentials, difficulty gap, damage type, power-spike timing, style matchup, and hard-counter rules.',
        'Engineered the original version as a multi-sheet Excel workbook containing over 100 interdependent hard-counter rules, translating complex, conditional game logic into a structured, maintainable spreadsheet model.',
        'Currently leading a full rebuild of the tool into a production web application, migrating the rules engine and data model from Excel/VBA into a Python and Supabase-backed architecture deployed via GitHub Pages.',
        'Applied systematic, rule-based thinking, carried over directly from structural QA/QC work, to design a scoring system that must remain internally consistent across hundreds of interacting variables.',
    ],
)

add_project_block(
    doc,
    title='Crypto Market Dashboard & ETL Pipeline',
    status='Live',
    stack='JavaScript, Python / yfinance, Supabase, GitHub Actions',
    links_text='Demo: burny143.github.io/crypto-etl/  |  Code: github.com/burny143/crypto-etl',
    bullets=[
        'Built and shipped a dark-themed, real-time cryptocurrency market dashboard featuring live candlestick charts and technical indicators, including Simple Moving Average (SMA) and Relative Strength Index (RSI).',
        'Developed a Python-based ETL pipeline using the yfinance library to extract OHLCV (Open, High, Low, Close, Volume) market data across three timeframes (daily, hourly, and 4-hour) for five separate cryptocurrencies.',
        'Automated the entire data pipeline end to end using GitHub Actions, scheduling recurring extraction jobs that upsert clean, de-duplicated market data directly into a Supabase database.',
        'Delivered a fully live, self-sustaining data product, from raw market data ingestion through to a polished, user-facing analytics interface, without a team or existing template to build from.',
    ],
)

# ══════════════════════════════════════════════════════════════
# WORK EXPERIENCE
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, 'Work Experience', level=2)

# --- INDRA ---
add_subheading(doc, 'Network & Deskside Support Engineer')
add_meta_line(doc, 'INDRA \u2014 Asian Development Bank (ADB) Project  |  Ortigas, Pasig City  |  May 2025 \u2013 Present')
for b in [
    'Resolved Level 2 network escalations by performing root-cause isolation on Zscaler connectivity failures, DNS conflicts, website whitelisting requests, and Wi-Fi/internet issues across ADB office and remote environments.',
    'Served as the primary point of contact for Zscaler Private Access (ZPA) firewall onboarding, coordinating access requirements directly between IT Security and the Network team.',
    'Owned incidents end to end, delivering clear, consistent status communication from initial isolation through resolution or escalation to Level 3 support.',
    'Conducted on-site diagnostics, including physical site visits to verify Wi-Fi coverage and wall port connectivity, and coordinated port activations directly with the Facilities team.',
    'Delivered deskside support covering hardware provisioning, Active Directory password resets, hardware diagnostics, and detailed incident logging in ServiceNow.',
]:
    add_bullet(doc, b)

# --- Accenture ---
add_subheading(doc, 'Custom Associate Software Engineer')
add_meta_line(doc, 'Accenture  |  Mandaluyong  |  May 2023 \u2013 Feb 2025')
for b in [
    'Managed Tier-2 support across a portfolio of more than 50 enterprise applications, spanning high-security production systems to niche, low-volume internal tools.',
    'Investigated and resolved production data errors using Oracle SQL and Microsoft SQL, tracing data flow issues through integration pipelines built on Dell Boomi and TIBCO.',
    'Authored and executed formal Change Requests and Service Requests to support backend updates and Windows Server / SQL Server upgrades, consistently achieving zero-downtime deployments.',
    'Served on a rotational basis as Incident Coordinator for P1 to P3 severity incidents, maintaining uninterrupted global operations throughout each incident lifecycle.',
    'Collaborated directly with client Subject Matter Experts (SMEs) and third-party vendors on bug diagnosis, root-cause analysis, and patch deployment.',
]:
    add_bullet(doc, b)

# --- Independent ---
add_subheading(doc, 'Independent \u2014 AI-Assisted Application Development')
add_meta_line(doc, 'Personal Projects  |  Self-directed  |  2025 \u2013 Present')
for b in [
    'Independently designed, built, and deployed a multi-sheet MLBB hero counter-pick scoring engine, evaluating matchups across role advantage, stat differentials, damage type, power-spike timing, and over 100 hard-counter rules; currently scaling the project from Excel into a full Supabase and Python-powered web application.',
    'Built a live cryptocurrency market dashboard with real-time candlestick charts and technical indicators, backed by a custom Python/yfinance data ingestion script and a fully automated GitHub Actions deployment pipeline into Supabase and GitHub Pages.',
    'Adopted AI as an active development collaborator across architecture decisions, debugging, and iterative refinement, while retaining full ownership of data structure design, formula and business logic, and end-to-end quality review.',
    'Self-directed ongoing study in AI engineering, covering prompt engineering, agentic workflows, and applied machine learning tooling.',
]:
    add_bullet(doc, b)

# --- Winsom ---
add_subheading(doc, 'Design and Engineering Operations In-Charge')
add_meta_line(doc, 'Winsom Building Products Co. Ltd.  |  Pasig / Ortigas  |  Jul 2021 \u2013 Sep 2022')
for b in [
    'Conducted site surveys and field measurements to support the creation of 2D CAD shop drawings, quantity take-offs, and cost estimates for architectural product proposals.',
    'Coordinated end-to-end material delivery logistics between client architects, warehouse teams, and site crews for specialized architectural building systems.',
    'Managed permitting and delivery scheduling in alignment with individual client site rules and requirements across multiple concurrent projects.',
    'Supervised third-party installation crews and performed quality inspections to verify compliance with approved shop drawings.',
]:
    add_bullet(doc, b)

# --- Hanston ---
add_subheading(doc, 'Project Civil Engineer')
add_meta_line(doc, 'Hanston Commercial and Industrial Corp.  |  Mandaluyong  |  Jan 2020 \u2013 Feb 2021')
for b in [
    'Acted as owner\u2019s site representative, performing daily structural QA/QC inspections on a multi-story commercial development to ensure strict adherence to approved plans, specifications, and building standards.',
    'Conducted formwork alignment and rebar inspections, formally signing off on pre-pour approvals before concrete casting operations.',
    'Coordinated with structural consultants and subcontractors to remediate honeycomb defects, including witnessing on-site high-strength repair mix testing.',
    'Supervised major concrete casting operations and delivered live progress monitoring and reporting to project owners.',
]:
    add_bullet(doc, b)

# --- Jetti ---
add_subheading(doc, 'Project Engineer / Project Manager')
add_meta_line(doc, 'Jetti Petroleum Inc.  |  Pasay City / Bulacan Sites  |  Nov 2018 \u2013 Jan 2020')
for b in [
    'Managed General Contractors through the full project lifecycle, from bidding through construction to final structural turnover, across multiple fuel station developments in Bulacan.',
    'Defined detailed project scopes and evaluated competing contractor proposals to select qualified bidders.',
    'Enforced material compliance against approved engineering drawings and site safety protocols throughout construction.',
    'Tracked project schedules using Gantt charts in MS Project, delivering consistent weekly progress reports to stakeholders.',
]:
    add_bullet(doc, b)

# --- CLG ---
add_subheading(doc, 'Project Supervisor / Project Manager')
add_meta_line(doc, 'CLG Shalom Builders  |  Bi\u00f1an, Laguna  |  Aug 2017 \u2013 Nov 2018')
for b in [
    'Managed off-hours interior renovation projects, including teller counter expansions, for BPI Southern Luzon bank branches within tight, weekend-only build windows.',
    'Met directly with BPI architects to draft 2D CAD layouts and secure formal client approvals ahead of site mobilization.',
    'Supervised two concurrent six-person site teams alongside specialized electrical subcontractors.',
    'Coordinated material sourcing, vehicle logistics, and weekly crew payroll across simultaneous project sites.',
]:
    add_bullet(doc, b)

# ══════════════════════════════════════════════════════════════
# EDUCATION
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, 'Education', level=2)

add_subheading(doc, 'Bachelor of Science in Civil Engineering')
add_meta_line(doc, 'Technological University of the Philippines \u2013 Taguig  |  2017 \u2013 2022')

add_subheading(doc, 'Secondary Education')
add_meta_line(doc, 'Taguig  |  2010 \u2013 2016')

# ══════════════════════════════════════════════════════════════
# CERTIFICATIONS & LICENSES
# ══════════════════════════════════════════════════════════════

add_heading_styled(doc, 'Certifications & Licenses', level=2)

certs = [
    ('Licensed Civil Engineer', 'Professional Regulation Commission (PRC), License No. 0147997 \u2014 Passed November 2016'),
    ('Dell Boomi Professional Integration Developer', 'Completed 2023'),
    ('Dell Boomi Associate Integration Developer', 'Completed 2023'),
    ('Self-Directed AI Engineering Study', 'Ongoing; prompt engineering, agentic workflows, and applied ML/AI tooling'),
]
for title, detail in certs:
    p = doc.add_paragraph()
    run_t = p.add_run(title)
    run_t.bold = True
    run_t.font.size = Pt(10.5)
    run_t.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)

    p2 = doc.add_paragraph()
    run_d = p2.add_run(detail)
    run_d.font.size = Pt(10)
    run_d.font.name = 'Calibri'
    run_d.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)

# ── Save ──
output_path = r'C:\Users\THINKPAD\OneDrive\Desktop\Fredrick\Latest CV - 2026\FREDRICK_MAHINAY_CV_2026.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'CV saved to: {output_path}')
